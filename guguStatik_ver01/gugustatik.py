import magic
import zipfile
import urllib.parse
import pandas as pd 
from docx import Document
from langdetect import detect
from oletools.olevba import VBA_Parser
import pefile
import math
import datetime
import re
import requests
import string
from collections import Counter
import re
from PyPDF2 import PdfReader
import rarfile
import pdfplumber
import py7zr
from alive_progress import alive_bar
from tkinter import Tk
from tkinter.filedialog import askopenfilename
from rich import print as rprint
from rich.table import Table
def print_greeting():
    print('''
                                                         .x+=:.        s                     s       .          ..      
                                                        z`    ^%      :8                    :8      @88>  < .z@8"`      
               x.    .                    x.    .          .   <k    .88                   .88      %8P    !@88E        
     uL      .@88k  z88u        uL      .@88k  z88u      .@8Ned8"   :888ooo       u       :888ooo    .     '888E   u    
 .ue888Nc.. ~"8888 ^8888    .ue888Nc.. ~"8888 ^8888    .@^%8888"  -*8888888    us888u.  -*8888888  .@88u    888E u@8NL  
d88E`"888E`   8888  888R   d88E`"888E`   8888  888R   x88:  `)8b.   8888    .@88 "8888"   8888    ''888E`   888E`"88*"  
888E  888E    8888  888R   888E  888E    8888  888R   8888N=*8888   8888    9888  9888    8888      888E    888E .dN.   
888E  888E    8888  888R   888E  888E    8888  888R    %8"    R88   8888    9888  9888    8888      888E    888E~8888   
888E  888E    8888 ,888B . 888E  888E    8888 ,888B .   @8Wou 9%   .8888Lu= 9888  9888   .8888Lu=   888E    888E '888&  
888& .888E   "8888Y 8888"  888& .888E   "8888Y 8888"  .888888P`    ^%888*   9888  9888   ^%888*     888&    888E  9888. 
*888" 888&    `Y"   'YP    *888" 888&    `Y"   'YP    `   ^"F        'Y"    "888*""888"    'Y"      R888" '"888*" 4888" 
 `"   "888E                 `"   "888E                                       ^Y"   ^Y'               ""      ""    ""   
.dWi   `88E                .dWi   `88E                                                                                  
4888~  J8%                 4888~  J8%                                                                                   
 ^"===*"`                   ^"===*"`                                                                                   
    ''')

def extract_info_from_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        text = '\n'.join(page.extract_text() for page in pdf.pages)

        # Write the extracted text to a .txt file
        with open('extracted_text.txt', 'w') as f:
            f.write(text)

    urls = re.findall(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', text)
    ips = re.findall(r'\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b', text)
    domains = [urllib.parse.urlparse(url).netloc for url in urls]

    return urls, ips, domains

def identify_file_type(file_path):
    mime = magic.Magic(mime=True)
    return mime.from_file(file_path)

def has_macros(filename):
    vbaparser = VBA_Parser(filename)
    return vbaparser.detect_vba_macros()

def is_password_protected(file_path, file_type):
    if "zip" in file_type:
        try:
            with zipfile.ZipFile(file_path) as zf:
                badfile = zf.testzip()
                if badfile:
                    return True
        except RuntimeError:
            return True
    elif "rar" in file_type:
        try:
            with rarfile.RarFile(file_path) as rf:
                rf.testrar()
        except rarfile.RarCannotExec:
            return True
    elif "7z" in file_type:
        try:
            with py7zr.SevenZipFile(file_path, mode='r') as z:
                z.testzip()
        except py7zr.exceptions.Bad7zFile:
            return True
        except py7zr.exceptions.PasswordRequired:
            return True
    
    elif "pdf" in file_type:
        try:
            with pdfplumber.open(file_path) as pdf:
                text = '\n'.join(page.extract_text() for page in pdf.pages)
            return False
        except:
            return True

def main_menu():
    while True:
        print("\n[1]: Select & Analyze a file")
        print("[2]: Exit")
        print("[3]: Analyze a PE or DLL")
        choice = input("Enter your choice: ")

        if choice == '1':
            filename = askopenfilename()
            analyze_file(filename)
        elif choice == '2':
            break
        elif choice == '3':
            filename = askopenfilename()
            pe_info = extract_info_from_pe(filename)

            table = Table(show_header=True, header_style="bold magenta")
            table.add_column("File Path")
            table.add_column("URLs")
            table.add_column("Domains")
            table.add_column("IP Addresses")
            table.add_column("Architecture")
            table.add_column("General Entropy")
            table.add_column("File Size")
            table.add_column("Number of Sections")
            table.add_column("Compilation Date")
            table.add_column("DLLs")
            table.add_column("Packed")
            table.add_column("Packing Algorithm")
            for i, section_info in enumerate(pe_info['sections_info'], start=1):
                section_table = Table(show_header=True, header_style="bold magenta")
                section_table.add_column("Section Name")
                section_table.add_column("Entropy")
                section_table.add_column("Virtual Size")
                section_table.add_column("Raw Size")
            
                row_data_section = [section_info['Name'], str(section_info['Entropy']), str(section_info['Virtual Size']), str(section_info['Raw Size'])]
                section_table.add_row(*row_data_section)
                rprint(section_table)

            row_data = [
                filename, 
                ", ".join(s.decode() for s in pe_info['urls']),  
                ", ".join(s.decode() for s in pe_info['domains']),  
                ", ".join(s.decode() for s in pe_info['ips']),  
                pe_info['architecture'], 
                str(pe_info['general_entropy']), 
                str(pe_info['file_size']), 
                str(pe_info['num_sections']), 
                str(pe_info['compilation_date']), 
                ", ".join(s.decode() for s in pe_info['dlls']),  
                str(pe_info['packed']), 
                pe_info['packing_algorithm']
            ]

            table.add_row(*row_data)
            rprint(table)
        else:
            print("Invalid choice. Please enter 1, 2 or 3.")
                       

def entropy(data):
    if not data:
        return 0.0
    occurences = Counter(bytearray(data))
    entropy = 0
    len_data = len(data)
    for x in occurences.values():
        p_x = float(x) / len_data
        entropy -= p_x * math.log(p_x, 2)
    return entropy

def get_tlds():
    tlds = requests.get('https://data.iana.org/TLD/tlds-alpha-by-domain.txt')
    return tlds.text.split('\n')[1:-1]

def analyze_sections(pe):
    sections_info = []
    packers_sections = {
        # The packer/protector/tools section names/keywords
        '.aspack': 'Aspack packer',
        '.adata': 'Aspack packer/Armadillo packer',
        'ASPack': 'Aspack packer',
        '.ASPack': 'ASPAck Protector',
        '.boom': 'The Boomerang List Builder (config+exe xored with a single byte key 0x77)',
        '.ccg': 'CCG Packer (Chinese Packer)',
        '.charmve': 'Added by the PIN tool',
        'BitArts': 'Crunch 2.0 Packer',
        'DAStub': 'DAStub Dragon Armor protector',
        '!EPack': 'Epack packer',
        'FSG!': 'FSG packer (not a section name, but a good identifier)',
        '.gentee': 'Gentee installer',
        'kkrunchy': 'kkrunchy Packer',
        '.mackt': 'ImpRec-created section',
        '.MaskPE': 'MaskPE Packer',
        'MEW': 'MEW packer',
        '.MPRESS1': 'Mpress Packer',
        '.MPRESS2': 'Mpress Packer',
        '.neolite': 'Neolite Packer',
        '.neolit': 'Neolite Packer',
        '.nsp1': 'NsPack packer',
        '.nsp0': 'NsPack packer',
        '.nsp2': 'NsPack packer',
        'nsp1': 'NsPack packer',
        'nsp0': 'NsPack packer',
        'nsp2': 'NsPack packer',
        '.packed': 'RLPack Packer (first section)',
        'pebundle': 'PEBundle Packer',
        'PEBundle': 'PEBundle Packer',
        'PEC2TO': 'PECompact packer',
        'PECompact2': 'PECompact packer (not a section name, but a good identifier)',
        'PEC2': 'PECompact packer',
        'pec1': 'PECompact packer',
        'pec2': 'PECompact packer',
        'PEC2MO': 'PECompact packer',
        'PELOCKnt': 'PELock Protector',
        '.perplex': 'Perplex PE-Protector',
        'PESHiELD': 'PEShield Packer',
        '.petite': 'Petite Packer',
        'petite': 'Petite Packer',
        '.pinclie': 'Added by the PIN tool',
        'ProCrypt': 'ProCrypt Packer',
        '.RLPack': 'RLPack Packer (second section)',
        '.rmnet': 'Ramnit virus marker',
        'RCryptor': 'RPCrypt Packer',
        '.RPCrypt': 'RPCrypt Packer',
        '.seau': 'SeauSFX Packer',
        '.sforce3': 'StarForce Protection',
        '.spack': 'Simple Pack (by bagie)',
        '.svkp': 'SVKP packer',
        'Themida': 'Themida Packer',
        '.Themida': 'Themida Packer',
        'Themida ': 'Themida Packer',
        '.taz': 'Some version os PESpin',
        '.tsuarch': 'TSULoader',
        '.tsustub': 'TSULoader',
        '.packed': 'Unknown Packer',
        'PEPACK!!': 'Pepack',
        '.Upack': 'Upack packer',
        '.ByDwing': 'Upack Packer',
        'UPX0': 'UPX packer',
        'UPX1': 'UPX packer',
        'UPX2': 'UPX packer',
        'UPX!': 'UPX packer',
        '.UPX0': 'UPX Packer',
        '.UPX1': 'UPX Packer',
        '.UPX2': 'UPX Packer',
        '.vmp0': 'VMProtect packer',
        '.vmp1': 'VMProtect packer',
        '.vmp2': 'VMProtect packer',
        'VProtect': 'Vprotect Packer',
        '.winapi': 'Added by API Override tool',
        'WinLicen': 'WinLicense (Themida) Protector',
        '_winzip_': 'WinZip Self-Extractor',
        '.WWPACK': 'WWPACK Packer',
        '.yP': 'Y0da Protector',
        '.y0da': 'Y0da Protector',
    }
   

    for section in pe.sections:
        section_name = section.Name.decode(errors='ignore')
        section_entropy = entropy(section.get_data())
        virtual_size = section.Misc_VirtualSize
        raw_size = len(section.get_data())

        packing_algorithm = packers_sections.get(section_name, 'Unknown')
        section_info = {
            "Name": section_name,
            "Entropy": section_entropy,
            "Virtual Size": virtual_size,
            "Raw Size": raw_size,
            "Packing Algorithm":packing_algorithm
        }

        sections_info.append(section_info)
    return sections_info



def extract_info_from_pe(filename):
    pe = pefile.PE(filename)

    # 3.1
    tlds = get_tlds()
    with open(filename, 'rb') as f:
        data = f.read()
    strings = re.findall(b'[A-Za-z0-9/\-:]{4,}', data)
    urls = [s for s in strings if b'http' in s]
    domains = [s for s in strings if any(tld.encode() in s for tld in tlds)]
    ips = [s for s in strings if re.match(b'\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}', s)]

    # 3.2
    architecture = 'x86' if pe.FILE_HEADER.Machine == 0x014c else 'x86-x64' if pe.FILE_HEADER.Machine == 0x8664 else 'Unknown'
    general_entropy = entropy(data)
    file_size = len(data)
    num_sections = len(pe.sections)
    sections = [(s.Name, entropy(s.get_data()), s.Misc_VirtualSize, len(s.get_data())) for s in pe.sections]
    compilation_date = datetime.datetime.fromtimestamp(pe.FILE_HEADER.TimeDateStamp)
    dlls = [entry.dll for entry in pe.DIRECTORY_ENTRY_IMPORT]

    # 3.3
    packed = any(s.get_entropy() > 7 for s in pe.sections)
    packing_algorithm = 'Unknown'
    sections_info = analyze_sections(pe)
    if packed:
        packing_algorithm = next((section['Packing Algorithm'] for section in sections_info if section['Packing Algorithm'] != 'Unknown'), "Unknown")
    else:
        packing_algorithm = 'Unknown'
    return {
        'urls': urls,
        'domains': domains,
        'ips': ips,
        'architecture': architecture,
        'general_entropy': general_entropy,
        'file_size': file_size,
        'num_sections': num_sections,
        'sections': sections,
        'compilation_date': compilation_date,
        'dlls': dlls,
        'packed': packed,
        'packing_algorithm': packing_algorithm,
        'sections_info': sections_info,
    }
def analyze_file(filename):
    with alive_bar(4, title='Analyzing...') as bar:
        file_type = identify_file_type(filename)
        bar()
        is_protected = is_password_protected(filename, file_type)
        bar()

        table = Table(show_header=True, header_style="bold magenta")
        table.add_column("File Path")
        table.add_column("File Type")
        table.add_column("Password Protected")

        row_data = [filename, file_type, "Yes" if is_protected else "No"]

        if "pdf" in file_type and not is_protected:
            urls, ips, domains = extract_info_from_pdf(filename)
            pdf_page_count = len(PdfReader(filename).pages)
            pdf_page_count= math.ceil(pdf_page_count)
            table.add_column("URLs")
            table.add_column("IP Addresses")
            table.add_column("Domain Names")
            table.add_column("Page Count")
            row_data.extend([", ".join(urls), ", ".join(ips), ", ".join(domains), (str(pdf_page_count))])
        elif "doc" in file_type and not is_protected:
            doc = Document(filename)
            text = ' '.join(paragraph.text for paragraph in doc.paragraphs)
            language = detect(text)
            num_pages = len(doc.paragraphs)  
            contains_macros = has_macros(filename)  

            table.add_column("Language")
            table.add_column("Number of Pages")
            table.add_column("Contains Macros")
            row_data.extend([language, str(num_pages), "Yes" if contains_macros else "No"])  
        bar()

        table.add_row(*row_data)
        bar()

    rprint(table)

print_greeting()
main_menu()