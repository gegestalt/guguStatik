from ensurepip import bootstrap
import os
from flask import Flask, redirect, render_template, request, url_for
from rich.console import Console
from rich.table import Table
import tempfile
from werkzeug.utils import secure_filename

import magic
import zipfile
import urllib.parse
from flask import render_template
import pandas as pd 
from docx import Document
from langdetect import detect
from oletools.olevba import VBA_Parser
import pefile
import math
import datetime
import re
import requests
import string
from collections import Counter
import re
from PyPDF2 import PdfReader
import rarfile
import pdfplumber
import py7zr
from alive_progress import alive_bar
from tkinter import Tk
from tkinter.filedialog import askopenfilename
from rich import print as rprint
from rich.table import Table
def print_greeting():
    print('''
                                                         .x+=:.        s                     s       .          ..      
                                                        z`    ^%      :8                    :8      @88>  < .z@8"`      
               x.    .                    x.    .          .   <k    .88                   .88      %8P    !@88E        
     uL      .@88k  z88u        uL      .@88k  z88u      .@8Ned8"   :888ooo       u       :888ooo    .     '888E   u    
 .ue888Nc.. ~"8888 ^8888    .ue888Nc.. ~"8888 ^8888    .@^%8888"  -*8888888    us888u.  -*8888888  .@88u    888E u@8NL  
d88E`"888E`   8888  888R   d88E`"888E`   8888  888R   x88:  `)8b.   8888    .@88 "8888"   8888    ''888E`   888E`"88*"  
888E  888E    8888  888R   888E  888E    8888  888R   8888N=*8888   8888    9888  9888    8888      888E    888E .dN.   
888E  888E    8888  888R   888E  888E    8888  888R    %8"    R88   8888    9888  9888    8888      888E    888E~8888   
888E  888E    8888 ,888B . 888E  888E    8888 ,888B .   @8Wou 9%   .8888Lu= 9888  9888   .8888Lu=   888E    888E '888&  
888& .888E   "8888Y 8888"  888& .888E   "8888Y 8888"  .888888P`    ^%888*   9888  9888   ^%888*     888&    888E  9888. 
*888" 888&    `Y"   'YP    *888" 888&    `Y"   'YP    `   ^"F        'Y"    "888*""888"    'Y"      R888" '"888*" 4888" 
 `"   "888E                 `"   "888E                                       ^Y"   ^Y'               ""      ""    ""   
.dWi   `88E                .dWi   `88E                                                                                  
4888~  J8%                 4888~  J8%                                                                                   
 ^"===*"`                   ^"===*"`                                                                                   
    ''')
app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'


if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.static_folder = 'static'
api_key = "a468dec43448580f46c2f928fa08cb52ab9397b0acdb090eaeadb3bfea6fa6e9"
def submit_file_to_virustotal(filename):
    url = "https://www.virustotal.com/api/v3/files"
    files = {"file": (filename, open(filename, "rb"), "application/x-msdownload")}
    headers = {"accept": "application/json", "x-apikey": api_key}

    response = requests.post(url, files=files, headers=headers)
    return response.json()

def get_analysis_results(analysis_id):
    url = f"https://www.virustotal.com/api/v3/analyses/{analysis_id}"
    headers = {"accept": "application/json", "x-apikey": api_key}

    response = requests.get(url, headers=headers)
    return response.json()



from werkzeug.utils import secure_filename

@app.route("/vt_analysis", methods=["POST"])
def submit_file():
    if request.method == "POST":
        file = request.files["file"]
        if file:
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            response = submit_file_to_virustotal(file_path)
            analysis_id = response["data"]["id"]
            return redirect(url_for('display_results', analysis_id=analysis_id))
        else:
            return "No file received", 400

@app.route("/results/<analysis_id>",methods=["GET"])
def display_results(analysis_id):
        response = get_analysis_results(analysis_id)
        results = response["data"]["attributes"]["results"]
        attributes = response["data"]["attributes"]
        meta = response.get("meta", {}) 
        return render_template("vt_report.html", results=results,meta=meta,attributes=attributes)

def extract_info_from_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        text = '\n'.join(page.extract_text() for page in pdf.pages)

        with open('extracted_text.txt', 'w') as f:
            f.write(text)

    urls = re.findall(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', text)
    ips = re.findall(r'\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b', text)
    domains = [urllib.parse.urlparse(url).netloc for url in urls]

    return urls, ips, domains

def identify_file_type(file_path):
    mime = magic.Magic(mime=True)
    return mime.from_file(file_path)

def has_macros(filename):
    vbaparser = VBA_Parser(filename)
    return vbaparser.detect_vba_macros()

def is_password_protected(file_path, file_type):
    if "zip" in file_type:
        try:
            with zipfile.ZipFile(file_path) as zf:
                badfile = zf.testzip()
                if badfile:
                    return True
        except RuntimeError:
            return True
    elif "rar" in file_type:
        try:
            with rarfile.RarFile(file_path) as rf:
                rf.testrar()
        except rarfile.RarCannotExec:
            return True
    elif "7z" in file_type:
        try:
            with py7zr.SevenZipFile(file_path, mode='r') as z:
                z.testzip()
        except py7zr.exceptions.Bad7zFile:
            return True
        except py7zr.exceptions.PasswordRequired:
            return True
    
    elif "pdf" in file_type:
        try:
            with pdfplumber.open(file_path) as pdf:
                text = '\n'.join(page.extract_text() for page in pdf.pages)
            return False
        except:
            return True

def main_menu():
    while True:
        print("\n[1]: Select & Analyze a file")
        print("[2]: Exit")
        print("[3]: Analyze a PE or DLL")
        choice = input("Enter your choice: ")

        if choice == '1':
            filename = askopenfilename()
            analyze_file(filename)
        elif choice == '2':
            break
        elif choice == '3':
            filename = askopenfilename()
            pe_info = extract_info_from_pe(filename)

            table = Table(show_header=True, header_style="bold magenta")
            table.add_column("File Path")
            table.add_column("URLs")
            table.add_column("Domains")
            table.add_column("IP Addresses")
            table.add_column("Architecture")
            table.add_column("General Entropy")
            table.add_column("File Size")
            table.add_column("Number of Sections")
            table.add_column("Compilation Date")
            table.add_column("DLLs")
            table.add_column("Packed")
            table.add_column("Packing Algorithm")
            for i, section_info in enumerate(pe_info['sections_info'], start=1):
                section_table = Table(show_header=True, header_style="bold magenta")
                section_table.add_column("Section Name")
                section_table.add_column("Entropy")
                section_table.add_column("Virtual Size")
                section_table.add_column("Raw Size")
            
                row_data_section = [section_info['Name'], str(section_info['Entropy']), str(section_info['Virtual Size']), str(section_info['Raw Size'])]
                section_table.add_row(*row_data_section)
                rprint(section_table)

            row_data = [
                filename, 
                ", ".join(s.decode() for s in pe_info['urls']),  
                ", ".join(s.decode() for s in pe_info['domains']),  
                ", ".join(s.decode() for s in pe_info['ips']),  
                pe_info['architecture'], 
                str(pe_info['general_entropy']), 
                str(pe_info['file_size']), 
                str(pe_info['num_sections']), 
                str(pe_info['compilation_date']), 
                ", ".join(s.decode() for s in pe_info['dlls']),  
                str(pe_info['packed']), 
                pe_info['packing_algorithm']
            ]

            table.add_row(*row_data)
            rprint(table)
        else:
            print("Invalid choice. Please enter 1, 2 or 3.")
                       

def entropy(data):
    if not data:
        return 0.0
    occurences = Counter(bytearray(data))
    entropy = 0
    len_data = len(data)
    for x in occurences.values():
        p_x = float(x) / len_data
        entropy -= p_x * math.log(p_x, 2)
    return entropy

def get_tlds():
    tlds = requests.get('https://data.iana.org/TLD/tlds-alpha-by-domain.txt')
    return tlds.text.split('\n')[1:-1]

def analyze_sections(pe):
    sections_info = []
    packers_sections = {
        '.aspack': 'Aspack packer',
        '.adata': 'Aspack packer/Armadillo packer',
        'ASPack': 'Aspack packer',
        '.ASPack': 'ASPAck Protector',
        '.boom': 'The Boomerang List Builder (config+exe xored with a single byte key 0x77)',
        '.ccg': 'CCG Packer (Chinese Packer)',
        '.charmve': 'Added by the PIN tool',
        'BitArts': 'Crunch 2.0 Packer',
        'DAStub': 'DAStub Dragon Armor protector',
        '!EPack': 'Epack packer',
        'FSG!': 'FSG packer (not a section name, but a good identifier)',
        '.gentee': 'Gentee installer',
        'kkrunchy': 'kkrunchy Packer',
        '.mackt': 'ImpRec-created section',
        '.MaskPE': 'MaskPE Packer',
        'MEW': 'MEW packer',
        '.MPRESS1': 'Mpress Packer',
        '.MPRESS2': 'Mpress Packer',
        '.neolite': 'Neolite Packer',
        '.neolit': 'Neolite Packer',
        '.nsp1': 'NsPack packer',
        '.nsp0': 'NsPack packer',
        '.nsp2': 'NsPack packer',
        'nsp1': 'NsPack packer',
        'nsp0': 'NsPack packer',
        'nsp2': 'NsPack packer',
        '.packed': 'RLPack Packer (first section)',
        'pebundle': 'PEBundle Packer',
        'PEBundle': 'PEBundle Packer',
        'PEC2TO': 'PECompact packer',
        'PECompact2': 'PECompact packer (not a section name, but a good identifier)',
        'PEC2': 'PECompact packer',
        'pec1': 'PECompact packer',
        'pec2': 'PECompact packer',
        'PEC2MO': 'PECompact packer',
        'PELOCKnt': 'PELock Protector',
        '.perplex': 'Perplex PE-Protector',
        'PESHiELD': 'PEShield Packer',
        '.petite': 'Petite Packer',
        'petite': 'Petite Packer',
        '.pinclie': 'Added by the PIN tool',
        'ProCrypt': 'ProCrypt Packer',
        '.RLPack': 'RLPack Packer (second section)',
        '.rmnet': 'Ramnit virus marker',
        'RCryptor': 'RPCrypt Packer',
        '.RPCrypt': 'RPCrypt Packer',
        '.seau': 'SeauSFX Packer',
        '.sforce3': 'StarForce Protection',
        '.spack': 'Simple Pack (by bagie)',
        '.svkp': 'SVKP packer',
        'Themida': 'Themida Packer',
        '.Themida': 'Themida Packer',
        'Themida ': 'Themida Packer',
        '.taz': 'Some version os PESpin',
        '.tsuarch': 'TSULoader',
        '.tsustub': 'TSULoader',
        '.packed': 'Unknown Packer',
        'PEPACK!!': 'Pepack',
        '.Upack': 'Upack packer',
        '.ByDwing': 'Upack Packer',
        'UPX0': 'UPX packer',
        'UPX1': 'UPX packer',
        'UPX2': 'UPX packer',
        'UPX!': 'UPX packer',
        '.UPX0': 'UPX Packer',
        '.UPX1': 'UPX Packer',
        '.UPX2': 'UPX Packer',
        '.vmp0': 'VMProtect packer',
        '.vmp1': 'VMProtect packer',
        '.vmp2': 'VMProtect packer',
        'VProtect': 'Vprotect Packer',
        '.winapi': 'Added by API Override tool',
        'WinLicen': 'WinLicense (Themida) Protector',
        '_winzip_': 'WinZip Self-Extractor',
        '.WWPACK': 'WWPACK Packer',
        '.yP': 'Y0da Protector',
        '.y0da': 'Y0da Protector',
    }
   

    for section in pe.sections:
        section_name = section.Name.decode(errors='ignore')
        section_entropy = entropy(section.get_data())
        virtual_size = section.Misc_VirtualSize
        raw_size = len(section.get_data())

        packing_algorithm = packers_sections.get(section_name, 'Unknown')
        section_info = {
            "Name": section_name,
            "Entropy": section_entropy,
            "Virtual Size": virtual_size,
            "Raw Size": raw_size,
            "Packing Algorithm":packing_algorithm
        }

        sections_info.append(section_info)
    return sections_info

def extract_strings_from_pe(pe):
    strings = []
    for section in pe.sections:
        section_data = section.get_data()
        section_strings = re.findall(b'[A-Za-z0-9/\-:]{4,}', section_data)
        strings.extend(section_strings)
    return strings

@app.route('/extract_info_from_pe', methods=['POST'])
def extract_info_from_pe():
    uploaded_file = request.files['file']
    with tempfile.NamedTemporaryFile(delete=False) as temp_file:
        uploaded_file.save(temp_file.name)
        pe = pefile.PE(temp_file.name)

        
        tlds = get_tlds()
        with open(temp_file.name, 'rb') as f:
            data = f.read()
        strings=extract_strings_from_pe(pe)
        urls = [s for s in strings if b'http' in s]
        domains = [s for s in strings if any(tld.encode() in s for tld in tlds)]
        ips = [s for s in strings if re.match(b'\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}', s)]

        architecture = 'x86' if pe.FILE_HEADER.Machine == 0x014c else 'x86-x64' if pe.FILE_HEADER.Machine == 0x8664 else 'Unknown'
        general_entropy = entropy(data)
        file_size = len(data)
        num_sections = len(pe.sections)
        sections = [(s.Name, entropy(s.get_data()), s.Misc_VirtualSize, len(s.get_data())) for s in pe.sections]
        compilation_date = datetime.datetime.fromtimestamp(pe.FILE_HEADER.TimeDateStamp)
        dlls = [entry.dll for entry in pe.DIRECTORY_ENTRY_IMPORT]

        packed = any(s.get_entropy() > 7 for s in pe.sections)
        packing_algorithm = 'Unknown'
        sections_info = analyze_sections(pe)
        if packed:
            packing_algorithm = next((section['Packing Algorithm'] for section in sections_info if section['Packing Algorithm'] != 'Unknown'), "Unknown")
        else:
            packing_algorithm = 'Unknown'
        dlls = [dll.decode('utf-8') for dll in dlls]

        table = Table(show_header=True, header_style="bold magenta")
        table.add_column("URLs")
        table.add_column("Domains")
        table.add_column("IPs")
        table.add_column("Architecture")
        table.add_column("General Entropy")
        table.add_column("File Size")
        table.add_column("Number of Sections")
        table.add_column("Compilation Date")
        table.add_column("DLLs")
        table.add_column("Packed")
        table.add_column("Packing Algorithm")
        
        row_data = [
            str(urls), 
            str(domains), 
            str(ips), 
            architecture, 
            str(general_entropy), 
            str(file_size), 
            str(num_sections), 
            str(compilation_date), 
            str(dlls), 
            str(packed), 
            packing_algorithm
        ]

        table.add_row(*row_data)
        
        sections_table = Table(show_header=True, header_style="bold magenta")
        sections_table.add_column("Section Name")
        sections_table.add_column("Entropy")
        sections_table.add_column("Virtual Size")
        sections_table.add_column("Raw Size")
        sections_table.add_column("Packing Algorithm")
        for section_info in sections_info:
            row_data_section = [section_info['Name'], str(section_info['Entropy']), str(section_info['Virtual Size']), str(section_info['Raw Size']), section_info['Packing Algorithm']]
            sections_table.add_row(*row_data_section)

    return render_template('pe_result.html', urls=urls, domains=domains, ips=ips, architecture=architecture, general_entropy=general_entropy,
                        file_size=file_size, num_sections=num_sections, compilation_date=compilation_date, dlls=dlls,
                        packed=packed, packing_algorithm=packing_algorithm, sections_info=sections_info, strings=strings)



def analyze_file(filename):
    with alive_bar(4, title='Analyzing...') as bar:
        file_type = identify_file_type(filename)
        bar()
        is_protected = is_password_protected(filename, file_type)
        bar()

        table = Table(show_header=True, header_style="bold magenta")
        table.add_column("File Path")
        table.add_column("File Type")
        table.add_column("Password Protected")

        row_data = [filename, file_type, "Yes" if is_protected else "No"]

        if "pdf" in file_type and not is_protected:
            urls, ips, domains = extract_info_from_pdf(filename)
            pdf_page_count = len(PdfReader(filename).pages)
            pdf_page_count = math.ceil(pdf_page_count)
            table.add_column("URLs")
            table.add_column("IP Addresses")
            table.add_column("Domain Names")
            table.add_column("Page Count")
            row_data.extend([", ".join(urls), ", ".join(ips), ", ".join(domains), (str(pdf_page_count))])
        elif "doc" or "docx" in file_type and not is_protected:
            doc = Document(filename)
            num_pages = 0
            text = ""
            prev_paragraph = None
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    if prev_paragraph and prev_paragraph.style.name != paragraph.style.name:
                        num_pages += 1
                    prev_paragraph = paragraph
                    text += paragraph.text

            language = detect(text)
            contains_macros = has_macros(filename)

            table.add_column("Language")
            table.add_column("Number of Pages")
            table.add_column("Contains Macros")
            row_data.extend([language, str(num_pages), "Yes" if contains_macros else "No"])  
        bar()

        table.add_row(*row_data)
        bar()

    rprint(table)


@app.route('/')
def index():
    return render_template('index.html')

@app.route('/analyze', methods=['POST'])
def analyze():
    uploaded_file = request.files['file']
    filename = tempfile.NamedTemporaryFile(delete=False)
    uploaded_file.save(filename.name)
    file_type = identify_file_type(filename.name)
    is_protected = is_password_protected(filename.name, file_type)

    if "pdf" in file_type and not is_protected:
        urls, ips, domains = extract_info_from_pdf(filename.name)
        pdf_page_count = len(PdfReader(filename.name).pages)
        pdf_page_count = math.ceil(pdf_page_count)
        analysis_result = {
            "File Path": filename.name,
            "File Type": file_type,
            "Password Protected": "Yes" if is_protected else "No",
            "URLs": ", ".join(urls) if urls else "Not Found",
            "IP Addresses": ", ".join(ips) if ips else "Not Found",
            "Domain Names": ", ".join(domains) if domains else "Not Found",
            "Page Count": pdf_page_count
        }
    elif "doc" in file_type and not is_protected:
        doc = Document(filename.name)
        text = ' '.join(paragraph.text for paragraph in doc.paragraphs)
        language = detect(text)
        num_pages = len(doc.paragraphs)  
        contains_macros = has_macros(filename.name)  
        analysis_result = {
            "File Path": filename.name,
            "File Type": file_type,
            "Password Protected": "Yes" if is_protected else "No",
            "Language": language,
            "Number of Pages": num_pages,
            "Contains Macros": "Yes" if contains_macros else "No"
        }
    else:
        analysis_result = {
            "File Path": filename.name,
            "File Type": file_type,
            "Password Protected": "Yes" if is_protected else "No",
        }



    return render_template('result.html', result=analysis_result)
@app.route('/readme')
def readme():
    return render_template('readme.html')
if __name__ == '__main__':
    app.run(debug=True)
